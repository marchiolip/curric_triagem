import os
import spacy
from PyPDF2 import PdfReader
from docx import Document
from difflib import SequenceMatcher
import json
from datetime import datetime
from werkzeug.datastructures import FileStorage

# Carregar o modelo de linguagem do spaCy
nlp = spacy.load("pt_core_news_sm")

# Caminho para a pasta de perfis
PERFIS_DIR = 'PERFIS'
DB_FILE = 'db.json'  # Define o arquivo para armazenar os resultados

def salvar_perfil(nome_perfil, perfil):
    if not os.path.exists(PERFIS_DIR):
        os.mkdir(PERFIS_DIR)
    perfil_path = os.path.join(PERFIS_DIR, f'{nome_perfil}.json')
    with open(perfil_path, 'w') as arquivo_json:
        json.dump(perfil, arquivo_json)

def listar_perfis_disponiveis():
    perfis = []
    if os.path.exists(PERFIS_DIR):
        for filename in os.listdir(PERFIS_DIR):
            if filename.endswith('.json'):
                perfis.append(filename[:-5])  # Remove a extensão .json
    return perfis

def carregar_perfil(nome_perfil):
    perfil_path = os.path.join(PERFIS_DIR, f'{nome_perfil}.json')
    if os.path.exists(perfil_path):
        with open(perfil_path, 'r') as arquivo_json:
            return json.load(arquivo_json)
    return {}

def extract_text(file_input):
    if isinstance(file_input, FileStorage):  # Verifica se o input é um objeto FileStorage
        if file_input.filename.lower().endswith('.pdf'):
            reader = PdfReader(file_input.stream)
            text = "".join(page.extract_text() or "" for page in reader.pages)
        elif file_input.filename.lower().endswith('.docx'):
            doc = Document(file_input.stream)
            text = "\n".join(para.text for para in doc.paragraphs if para.text)
        else:
            raise ValueError("Formato de arquivo não suportado")
    else:
        if file_input.lower().endswith('.pdf'):
            reader = PdfReader(file_input)
            text = "".join(page.extract_text() or "" for page in reader.pages)
        elif file_input.lower().endswith('.docx'):
            doc = Document(file_input)
            text = "\n".join(para.text for para in doc.paragraphs if para.text)
        else:
            raise ValueError("Formato de arquivo não suportado")
    return text

def preprocess(text):
    doc = nlp(text)
    return [token.lemma_.lower() for token in doc if not token.is_stop and not token.is_punct and not token.is_space]

def calcular_pontuacao(curriculo_tokens, descricao_tokens, pesos_palavras_chave):
    if not pesos_palavras_chave:
        print("Não há pesos definidos no perfil.")
        return 0

    score = sum(pesos_palavras_chave.get(token, 0) * curriculo_tokens.count(token) for token in set(curriculo_tokens))
    similarity_score = calculate_similarity(' '.join(curriculo_tokens), ' '.join(descricao_tokens))
    final_score = score * similarity_score
    
    return final_score

def calculate_similarity(text1, text2):
    return SequenceMatcher(None, text1, text2).ratio()

def registrar_resultado(perfil_selecionado, nome_curriculo, pontuacao):
    # Carrega o db.json atual ou cria um novo se não existir
    db_data = []
    if os.path.exists(DB_FILE):
        with open(DB_FILE, 'r') as arquivo_db:
            db_data = json.load(arquivo_db)
    
    # Adiciona a data de análise
    data_analisado = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Verifica se o resultado já existe no banco de dados
    resultado_existente = any(
        resultado['perfil_selecionado'] == perfil_selecionado and
        resultado['nome_curriculo'] == nome_curriculo
        for resultado in db_data
    )
    
    # Se não existir e a pontuação for maior ou igual a 0.5, adiciona o novo resultado
    if not resultado_existente and pontuacao >= 0.5:
        db_data.append({
            'perfil_selecionado': perfil_selecionado,
            'nome_curriculo': nome_curriculo,
            'pontuacao': pontuacao,
            'data_analisado': data_analisado
        })

        # Salva o banco de dados atualizado
        with open(DB_FILE, 'w') as arquivo_db:
            json.dump(db_data, arquivo_db, indent=4)

